"""Report generation utilities."""

from __future__ import annotations

import io
from datetime import datetime


def _correlation_label(value: float) -> str:
    abs_val = abs(value)
    direction = "positiva" if value > 0 else "negativa"
    if abs_val >= 0.7:
        strength = "fuerte"
    elif abs_val >= 0.4:
        strength = "moderada"
    else:
        strength = "debil"
    return f"{strength} {direction}"


def generate_executive_summary(report, model_result=None) -> str:
    """Generate markdown summary for display in Streamlit."""
    lines = [
        "## Resumen ejecutivo", "",
        f"**Variable objetivo analizada:** `{report.target}`",
        f"**Variables predictoras seleccionadas:** {len(report.features)}", "",
        "### Correlaciones con la variable objetivo",
    ]

    if report.top_correlations.empty:
        lines.append("No se pudieron calcular correlaciones.")
    else:
        for _, row in report.top_correlations.head(3).iterrows():
            strength = _correlation_label(row["correlation"])
            lines.append(
                f"- **{row['feature']}**: correlacion {strength} ({row['correlation']:.2f})"
            )
        lines.append("")
        lines.append(
            "> Correlacion no implica causalidad. Estas relaciones son descriptivas, no explicativas."
        )

    if model_result is not None:
        lines += ["", "### Resultados del modelo de regresion lineal"]

        if model_result.excluded_identifiers:
            lines.append(
                f"> Variables excluidas por ser identificadores: "
                f"`{'`, `'.join(model_result.excluded_identifiers)}`. "
                f"No aportan valor predictivo real."
            )

        if model_result.leakage_warnings:
            for feat, corr in model_result.leakage_warnings:
                lines.append(
                    f"> **Alerta de Target Leakage:** `{feat}` "
                    f"tiene correlacion {corr} con el target. "
                    f"Considera excluirla del modelo."
                )

        lines += [
            f"- **MAE:** {model_result.mae} — error promedio absoluto",
            f"- **RMSE:** {model_result.rmse} — penaliza errores grandes",
            f"- **R2:** {model_result.r2:.4f} — varianza explicada en el conjunto de prueba",
            f"- **R2 ajustado:** {model_result.r2_adjusted} — penaliza predictoras innecesarias",
            f"- **Train:** {model_result.n_train} filas | **Test:** {model_result.n_test} filas",
        ]

        if model_result.small_dataset:
            lines.append(
                f"- **Validacion cruzada (k-fold) R2:** "
                f"{model_result.cv_r2_mean} +/- {model_result.cv_r2_std} "
                f"— mas confiable que el R2 simple para datasets pequenos"
            )

        lines.append("")

        if model_result.low_performance:
            lines.append(
                "> R2 < 0.30 — bajo desempeno. Considera revisar variables o usar otro modelo."
            )
        else:
            lines.append("> Desempeno aceptable para regresion lineal basica.")

        lines += [
            "", "### Limitaciones del modelo",
            "- Solo se usaron variables numericas como predictoras.",
            "- El modelo asume relaciones lineales entre variables.",
            "- Dataset pequeno puede no generalizar a datos nuevos.",
            "- Variables categoricas fueron ignoradas.",
            "- No se aplico regularizacion (Ridge/Lasso) ni seleccion de variables.",
        ]

    return "\n".join(lines)


def generate_txt_report(
    filename: str,
    profile,
    audit,
    report=None,
    model_result=None,
) -> str:
    """Generate a plain text report with all sections."""
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    sep = "=" * 60
    lines = [
        sep,
        "TABULAR INSIGHT WORKBENCH — REPORTE DE ANALISIS",
        sep,
        f"Archivo analizado : {filename}",
        f"Fecha de generacion: {now}",
        "",
        sep,
        "1. PERFIL DEL DATASET",
        sep,
        f"Filas             : {profile.rows:,}",
        f"Columnas          : {profile.columns:,}",
        f"Filas duplicadas  : {profile.duplicate_rows:,}",
        f"Celdas nulas      : {profile.missing_cells:,}",
        f"Porcentaje nulos  : {profile.missing_cells_pct:.2f}%",
    ]

    if profile.constant_columns:
        lines.append(f"Columnas constantes: {', '.join(profile.constant_columns)}")
    if profile.high_cardinality_columns:
        lines.append(f"Alta cardinalidad  : {', '.join(profile.high_cardinality_columns)}")

    lines += [
        "",
        sep,
        "2. AUDITORIA DE LIMPIEZA",
        sep,
        f"Filas antes       : {audit.rows_before:,}",
        f"Filas despues     : {audit.rows_after:,}",
        f"Columnas antes    : {audit.columns_before:,}",
        f"Columnas despues  : {audit.columns_after:,}",
        f"Duplicados elim.  : {audit.duplicates_removed}",
        f"Filas elim. nulos : {audit.rows_removed_missing}",
        f"Valores imputados : {audit.missing_values_filled}",
        f"Columnas eliminadas: {', '.join(audit.columns_removed) or 'Ninguna'}",
        f"Acciones aplicadas: {', '.join(audit.actions)}",
    ]

    if report is not None:
        lines += [
            "",
            sep,
            "3. ANALISIS EXPLORATORIO",
            sep,
            f"Variable objetivo : {report.target}",
            f"Predictoras       : {', '.join(report.features)}",
            "",
            "Correlaciones con la variable objetivo:",
        ]

        if report.top_correlations.empty:
            lines.append("  No se pudieron calcular correlaciones.")
        else:
            for _, row in report.top_correlations.head(5).iterrows():
                strength = _correlation_label(row["correlation"])
                lines.append(
                    f"  {row['feature']:<25} {row['correlation']:+.4f}  ({strength})"
                )

        lines.append("")
        lines.append("NOTA: Correlacion no implica causalidad.")

    if model_result is not None:
        lines += [
            "",
            sep,
            "4. MODELO DE REGRESION LINEAL",
            sep,
            f"Variable objetivo : {model_result.target}",
            f"Predictoras usadas: {', '.join(model_result.features)}",
            f"Filas entrenamiento: {model_result.n_train}",
            f"Filas prueba      : {model_result.n_test}",
            "",
            "METRICAS:",
            f"  MAE              : {model_result.mae}",
            f"  RMSE             : {model_result.rmse}",
            f"  R2               : {model_result.r2:.4f}",
            f"  R2 ajustado      : {model_result.r2_adjusted}",
        ]

        if model_result.small_dataset:
            lines.append(
                f"  CV R2 (k-fold)   : {model_result.cv_r2_mean} +/- {model_result.cv_r2_std}"
            )

        if model_result.excluded_identifiers:
            lines.append(
                f"\nVARIABLES EXCLUIDAS (identificadores): "
                f"{', '.join(model_result.excluded_identifiers)}"
            )

        if model_result.leakage_warnings:
            lines.append("\nALERTAS DE TARGET LEAKAGE:")
            for feat, corr in model_result.leakage_warnings:
                lines.append(f"  {feat}: correlacion {corr} con el target")

        if model_result.low_performance:
            lines.append("\nADVERTENCIA: R2 < 0.30 — bajo desempeno del modelo.")
        else:
            lines.append("\nDesempeno aceptable para regresion lineal basica.")

        lines += [
            "",
            "LIMITACIONES:",
            "  - Solo variables numericas fueron usadas como predictoras.",
            "  - El modelo asume relaciones lineales.",
            "  - Dataset pequeno puede no generalizar.",
            "  - Variables categoricas fueron ignoradas.",
            "  - Sin regularizacion ni seleccion automatica de variables.",
        ]

    lines += ["", sep, "Generado por Tabular Insight Workbench", sep]
    return "\n".join(lines)


def generate_docx_report(
    filename: str,
    profile,
    audit,
    report=None,
    model_result=None,
) -> bytes:
    """Generate a formatted Word (.docx) report and return as bytes."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title = doc.add_heading("Tabular Insight Workbench", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_heading("Reporte de Analisis de Datos", level=1)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    meta = doc.add_paragraph()
    meta.add_run("Archivo: ").bold = True
    meta.add_run(filename)
    meta.add_run("     Fecha: ").bold = True
    meta.add_run(now)

    doc.add_paragraph()

    # Section 1 — Dataset profile
    doc.add_heading("1. Perfil del Dataset", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Metrica"
    hdr[1].text = "Valor"
    for label, value in [
        ("Filas", f"{profile.rows:,}"),
        ("Columnas", f"{profile.columns:,}"),
        ("Filas duplicadas", f"{profile.duplicate_rows:,}"),
        ("Celdas nulas", f"{profile.missing_cells:,}"),
        ("Porcentaje nulos", f"{profile.missing_cells_pct:.2f}%"),
        ("Columnas constantes", ", ".join(profile.constant_columns) or "Ninguna"),
        ("Alta cardinalidad", ", ".join(profile.high_cardinality_columns) or "Ninguna"),
    ]:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(value)

    doc.add_paragraph()

    # Section 2 — Cleaning audit
    doc.add_heading("2. Auditoria de Limpieza", level=2)
    table2 = doc.add_table(rows=1, cols=2)
    table2.style = "Table Grid"
    hdr2 = table2.rows[0].cells
    hdr2[0].text = "Metrica"
    hdr2[1].text = "Valor"
    for label, value in [
        ("Filas antes", f"{audit.rows_before:,}"),
        ("Filas despues", f"{audit.rows_after:,}"),
        ("Columnas antes", f"{audit.columns_before:,}"),
        ("Columnas despues", f"{audit.columns_after:,}"),
        ("Duplicados eliminados", str(audit.duplicates_removed)),
        ("Filas eliminadas por nulos", str(audit.rows_removed_missing)),
        ("Valores imputados", str(audit.missing_values_filled)),
        ("Columnas eliminadas", ", ".join(audit.columns_removed) or "Ninguna"),
        ("Acciones aplicadas", ", ".join(audit.actions)),
    ]:
        row = table2.add_row().cells
        row[0].text = label
        row[1].text = str(value)

    doc.add_paragraph()

    # Section 3 — EDA
    if report is not None:
        doc.add_heading("3. Analisis Exploratorio", level=2)
        p = doc.add_paragraph()
        p.add_run("Variable objetivo: ").bold = True
        p.add_run(report.target)
        p2 = doc.add_paragraph()
        p2.add_run("Predictoras: ").bold = True
        p2.add_run(", ".join(report.features))

        doc.add_paragraph()
        doc.add_heading("Correlaciones con la variable objetivo", level=3)

        if not report.top_correlations.empty:
            table3 = doc.add_table(rows=1, cols=3)
            table3.style = "Table Grid"
            hdr3 = table3.rows[0].cells
            hdr3[0].text = "Variable"
            hdr3[1].text = "Correlacion"
            hdr3[2].text = "Intensidad"
            for _, row in report.top_correlations.head(5).iterrows():
                r = table3.add_row().cells
                r[0].text = str(row["feature"])
                r[1].text = f"{row['correlation']:+.4f}"
                r[2].text = _correlation_label(row["correlation"])

        doc.add_paragraph()
        note = doc.add_paragraph()
        note.add_run("Nota: ").bold = True
        note.add_run("Correlacion no implica causalidad. Estas relaciones son descriptivas.")

        doc.add_paragraph()

    # Section 4 — Model
    if model_result is not None:
        doc.add_heading("4. Modelo de Regresion Lineal", level=2)

        if model_result.excluded_identifiers:
            p = doc.add_paragraph()
            p.add_run("Variables excluidas (identificadores): ").bold = True
            p.add_run(", ".join(model_result.excluded_identifiers))

        if model_result.leakage_warnings:
            for feat, corr in model_result.leakage_warnings:
                p = doc.add_paragraph()
                run = p.add_run(f"ALERTA Target Leakage: {feat} correlaciona {corr} con el target.")
                run.bold = True

        doc.add_paragraph()
        table4 = doc.add_table(rows=1, cols=2)
        table4.style = "Table Grid"
        hdr4 = table4.rows[0].cells
        hdr4[0].text = "Metrica"
        hdr4[1].text = "Valor"

        metrics = [
            ("Variable objetivo", model_result.target),
            ("Predictoras usadas", ", ".join(model_result.features)),
            ("Filas entrenamiento", str(model_result.n_train)),
            ("Filas prueba", str(model_result.n_test)),
            ("MAE", str(model_result.mae)),
            ("RMSE", str(model_result.rmse)),
            ("R2", f"{model_result.r2:.4f}"),
            ("R2 ajustado", str(model_result.r2_adjusted)),
        ]

        if model_result.small_dataset:
            metrics.append((
                "CV R2 (k-fold)",
                f"{model_result.cv_r2_mean} +/- {model_result.cv_r2_std}"
            ))

        for label, value in metrics:
            r = table4.add_row().cells
            r[0].text = label
            r[1].text = str(value)

        doc.add_paragraph()

        verdict = doc.add_paragraph()
        if model_result.low_performance:
            verdict.add_run("ADVERTENCIA: R2 < 0.30 — bajo desempeno del modelo.").bold = True
        else:
            verdict.add_run("Desempeno aceptable para regresion lineal basica.").bold = True

        doc.add_paragraph()
        doc.add_heading("Limitaciones del modelo", level=3)
        for item in [
            "Solo variables numericas fueron usadas como predictoras.",
            "El modelo asume relaciones lineales entre variables.",
            "Dataset pequeno puede no generalizar a datos nuevos.",
            "Variables categoricas fueron ignoradas.",
            "Sin regularizacion (Ridge/Lasso) ni seleccion automatica de variables.",
        ]:
            doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Generado por Tabular Insight Workbench").italic = True

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()